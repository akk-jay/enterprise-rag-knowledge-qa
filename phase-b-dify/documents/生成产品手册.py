# -*- coding: utf-8 -*-
"""生成《产品手册-云办公平台》DOCX — CloudWork 产品手册 V3.5"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, datetime


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
FONT_FAMILY = "微软雅黑"
BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x1D, 0x1D, 0x1F)
GREY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xE8, 0xF0, 0xFE)


def _set_run(run, size=Pt(10.5), bold=False, color=DARK, font=FONT_FAMILY):
    run.font.size = size
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(8)
    for r in h.runs:
        r.font.name = FONT_FAMILY
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)
        r.font.color.rgb = BLUE
    return h


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    _set_run(run)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    _set_run(run)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p


def add_qa(doc, question, answer):
    """Add a Q&A block."""
    # Q
    qp = doc.add_paragraph()
    qp.paragraph_format.space_before = Pt(14)
    qp.paragraph_format.space_after = Pt(2)
    qr = qp.add_run(f"Q: {question}")
    _set_run(qr, size=Pt(11), bold=True, color=BLUE)
    # A
    ap = doc.add_paragraph()
    ap.paragraph_format.line_spacing = 1.5
    ap.paragraph_format.space_after = Pt(10)
    ar = ap.add_run(f"A: {answer}")
    _set_run(ar, size=Pt(10.5))
    return qp, ap


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


# ---------------------------------------------------------------------------
# main generator
# ---------------------------------------------------------------------------
def generate():
    doc = Document()

    # -- page setup ----------------------------------------------------------
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # ========================================================================
    # COVER PAGE
    # ========================================================================
    # blank vertical space
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # main title
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run("云办公平台")
    _set_run(r1, size=Pt(38), bold=True, color=BLUE)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("CloudWork")
    _set_run(r2, size=Pt(26), bold=False, color=RGBColor(0x66, 0x66, 0x66))

    # separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sep.add_run("━" * 30)
    _set_run(sr, size=Pt(10), color=BLUE)

    # subtitle
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st_run = st.add_run("产 品 手 册")
    _set_run(st_run, size=Pt(22), bold=True, color=DARK)

    # meta info
    for _ in range(3):
        doc.add_paragraph()

    meta_lines = [
        ("版本：", "V3.5"),
        ("更新日期：", "2025年6月"),
        ("文档密级：", "内部公开"),
    ]
    for label, value in meta_lines:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ml = mp.add_run(label)
        _set_run(ml, size=Pt(11), color=GREY)
        mv = mp.add_run(value)
        _set_run(mv, size=Pt(11), bold=True, color=DARK)

    # footer on cover
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(40)
    fr = fp.add_run("XX科技有限公司  |  cloudwork.example.com  |  400-XXX-XXXX")
    _set_run(fr, size=Pt(9), color=GREY)

    doc.add_page_break()

    # ========================================================================
    # TABLE OF CONTENTS
    # ========================================================================
    add_heading_styled(doc, "目  录", level=1)

    toc_items = [
        ("一、产品概述", 3),
        ("    1.1  产品简介", 3),
        ("    1.2  技术架构与安全", 3),
        ("二、核心功能", 5),
        ("    2.1  即时通讯", 5),
        ("    2.2  视频会议", 5),
        ("    2.3  任务管理", 5),
        ("    2.4  文档协作", 6),
        ("    2.5  日程管理", 6),
        ("三、快速入门", 7),
        ("    3.1  创建团队", 7),
        ("    3.2  邀请成员", 7),
        ("    3.3  创建项目", 8),
        ("四、常见问题", 9),
    ]

    for text, page in toc_items:
        tp = doc.add_paragraph()
        tp.paragraph_format.line_spacing = 1.8
        tp.paragraph_format.space_after = Pt(0)
        dots = "·" * (50 - len(text) * 2)
        tr = tp.add_run(f"{text} {dots} {page}")
        _set_run(tr, size=Pt(10.5))
        if not text.startswith("    "):
            tr.bold = True

    doc.add_page_break()

    # ========================================================================
    # 一、产品概述
    # ========================================================================
    add_heading_styled(doc, "一、产品概述", level=1)

    add_heading_styled(doc, "1.1  产品简介", level=2)
    add_body(
        doc,
        "CloudWork（云办公平台）是XX科技自主研发的新一代企业级协同办公平台，"
        "致力于为各类组织提供安全、高效、智能的一体化办公解决方案。平台深度整合了"
        "即时通讯（IM）、高清视频会议、任务管理、在线文档协作和智能日程管理五大核心能力，"
        "覆盖企业日常办公全场景需求。",
        indent=True,
    )
    add_body(
        doc,
        "CloudWork支持Windows、macOS、Linux三大桌面操作系统，以及iOS、Android移动端和"
        "Web浏览器访问，实现跨平台无缝衔接。自发布以来，平台已累计服务超过5000家企业客户，"
        "日活跃用户（DAU）突破50万，广泛应用于互联网、金融、制造、教育、医疗等行业。",
        indent=True,
    )

    add_heading_styled(doc, "1.2  技术架构与安全", level=2)
    add_body(
        doc,
        "CloudWork采用云原生微服务架构，基于Kubernetes容器编排实现弹性伸缩与高可用部署。"
        "平台支持SaaS公有云与私有化部署两种模式，满足不同规模企业的IT管理需求。"
        "私有化部署支持单机、集群两地三中心架构，保障业务连续性。",
        indent=True,
    )
    add_body(
        doc,
        "数据安全方面，CloudWork全链路采用AES-256加密传输与存储，通过ISO 27001信息安全管理体系认证"
        "及中国网络安全等级保护三级（等保三级）认证。平台定期进行第三方渗透测试与安全审计，"
        "确保客户数据的机密性、完整性与可用性。",
        indent=True,
    )

    # feature summary table
    add_body(doc, "平台核心指标一览：")
    table = doc.add_table(rows=7, cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_data = [
        ("支持平台", "Windows / macOS / Linux / iOS / Android / Web"),
        ("企业客户数", "5,000+"),
        ("日活跃用户 (DAU)", "500,000+"),
        ("视频会议并发", "单场最高 500 人"),
        ("消息加密", "端到端 AES-256"),
        ("认证资质", "ISO 27001 / 等保三级"),
        ("部署模式", "SaaS 公有云 + 私有化部署"),
    ]
    for i, (k, v) in enumerate(tbl_data):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.text = ""
        c1.text = ""
        k_run = c0.paragraphs[0].add_run(k)
        _set_run(k_run, size=Pt(10), bold=True)
        v_run = c1.paragraphs[0].add_run(v)
        _set_run(v_run, size=Pt(10))
        if i == 0:
            set_cell_shading(c0, "1A56DB")
            set_cell_shading(c1, "1A56DB")
            k_run.font.color.rgb = WHITE
            v_run.font.color.rgb = WHITE

    doc.add_page_break()

    # ========================================================================
    # 二、核心功能
    # ========================================================================
    add_heading_styled(doc, "二、核心功能", level=1)

    # 2.1 即时通讯
    add_heading_styled(doc, "2.1  即时通讯 (IM)", level=2)
    add_body(
        doc,
        "CloudWork即时通讯模块提供企业级安全沟通能力，支持单聊与群聊两种模式。"
        "群聊人数上限可达5000人，满足大型组织全员群的需求。消息类型涵盖文本、图片、"
        "语音、视频、文件（单文件最大2GB）及代码片段（支持30+编程语言语法高亮）。",
        indent=True,
    )
    add_bullet(doc, "消息撤回：发送后5分钟内可撤回，撤回后双方均显示撤回记录")
    add_bullet(doc, "阅读回执：支持已读/未读状态查看，群聊中可查看消息已读人员明细")
    add_bullet(doc, "端到端加密：敏感会话支持端到端加密模式，确保消息仅收发双方可解密")
    add_bullet(doc, "消息漫游：全平台消息历史云端同步，新设备登录后可回溯全部历史消息")

    # 2.2 视频会议
    add_heading_styled(doc, "2.2  视频会议", level=2)
    add_body(
        doc,
        "视频会议模块支持单场最高500人同时参会，提供高清1080P音视频传输。"
        "内置屏幕共享（支持全屏/窗口/标签页三种模式）、互动白板、虚拟背景、"
        "美颜滤镜等功能。会议录制文件自动保存至云端，默认保留30天，支持下载至本地。",
        indent=True,
    )
    add_bullet(doc, "实时字幕：基于AI语音识别，支持中英双语实时字幕，准确率超过97%")
    add_bullet(doc, "AI会议纪要：会议结束后5分钟内自动生成结构化纪要，包括议题摘要、待办事项与决议")
    add_bullet(doc, "等候室与入会密码：支持会前安全管控，防止未授权人员进入")
    add_bullet(doc, "分组讨论：主持人可将参会者分配至多个分组讨论室，结束后统一回归主会场")

    # 2.3 任务管理
    add_heading_styled(doc, "2.3  任务管理", level=2)
    add_body(
        doc,
        "任务管理模块采用项目-任务-子任务三级结构，支持看板、列表、甘特图、时间线"
        "四种视图模式，满足不同角色和场景的管理需求。用户可在任务间设置依赖关系，"
        "系统自动检测关键路径并标识瓶颈任务。",
        indent=True,
    )
    add_bullet(doc, "依赖管理：支持FS/FF/SS/SF四种依赖类型，自动计算关键路径与项目工期")
    add_bullet(doc, "智能预警：任务到期前3天自动推送提醒至负责人，逾期任务每日汇总通知")
    add_bullet(doc, "Excel导入导出：支持通过Excel模板批量导入/导出任务，适配企业已有项目数据")
    add_bullet(doc, "自定义字段：支持文本、数字、日期、下拉选项等自定义字段，灵活适配业务需求")

    # 2.4 文档协作
    add_heading_styled(doc, "2.4  文档协作", level=2)
    add_body(
        doc,
        "文档协作模块支持多人实时协同编辑，基于OT（Operational Transformation）算法"
        "实现毫秒级同步，避免编辑冲突。文档支持富文本编辑，包括表格、图片、代码块、"
        "数学公式（LaTeX）等多种内容形式。",
        indent=True,
    )
    add_bullet(doc, "编辑历史：完整记录每次编辑操作，支持任意历史版本预览与一键回滚")
    add_bullet(doc, "评论与协作：支持@提及成员添加评论，被@成员实时收到通知并可原地回复")
    add_bullet(doc, "权限管控：三级权限体系——查看、评论、编辑，按成员或部门精细授权")
    add_bullet(doc, "外部分享：支持生成分享链接，可设置访问密码与有效期，保障信息安全")

    # 2.5 日程管理
    add_heading_styled(doc, "2.5  日程管理", level=2)
    add_body(
        doc,
        "日程管理模块提供个人日历与团队日历双视图，支持日/周/月多种时间粒度展示。"
        "日程提醒支持多级时间配置（5分钟、15分钟、30分钟、1小时、1天前），"
        "提醒方式包括应用内推送、短信和邮件。",
        indent=True,
    )
    add_bullet(doc, "忙闲状态：团队成员间可查看忙闲状态，快速找到合适的会议时间")
    add_bullet(doc, "外部日历订阅：支持订阅Google Calendar、Outlook Calendar等外部日历，统一管理")
    add_bullet(doc, "语音创建：移动端支持语音输入创建日程，AI自动解析时间、地点与主题")
    add_bullet(doc, "周期日程：支持按天/周/月/年创建重复日程，灵活设置结束条件")

    doc.add_page_break()

    # ========================================================================
    # 三、快速入门
    # ========================================================================
    add_heading_styled(doc, "三、快速入门", level=1)
    add_body(doc, "本章将引导您完成CloudWork的基础设置，帮助团队快速上手使用。", indent=True)

    # 3.1 创建团队
    add_heading_styled(doc, "3.1  创建团队", level=2)
    add_body(
        doc,
        "注册CloudWork账号后，首先需要创建您的第一个团队。点击主界面左侧导航栏底部的"
        "\"+\"按钮，选择\"创建团队\"，进入创建向导。",
        indent=True,
    )
    add_body(doc, "创建团队的步骤如下：")
    add_bullet(doc, "第1步：输入团队名称（必填，支持中英文及特殊字符，2-30个字符）")
    add_bullet(doc, "第2步：选择团队类型——企业/政府/学校/其他，系统将据此配置默认模板")
    add_bullet(doc, "第3步：上传团队头像（可选，建议尺寸200x200px，支持JPG/PNG格式）")
    add_bullet(doc, "第4步：阅读并同意用户协议，点击\"完成创建\"")
    add_body(doc, "注意：每个账号最多可创建或加入5个团队。创建团队的用户自动成为团队超级管理员。")

    # 3.2 邀请成员
    add_heading_styled(doc, "3.2  邀请成员", level=2)
    add_body(doc, "CloudWork提供四种成员邀请方式，满足不同场景需求：", indent=True)
    add_bullet(doc, "通讯录搜索：在企业通讯录中直接搜索同事姓名或工号，一键添加至团队")
    add_bullet(doc, "邀请链接：生成团队专属邀请链接，通过微信/钉钉/邮件等渠道发送，受邀者点击即可加入")
    add_bullet(doc, "二维码邀请：生成团队二维码，受邀者使用CloudWork扫码即可加入（有效期24小时）")
    add_bullet(
        doc,
        "Excel批量导入：下载成员导入模板，按格式填写姓名、手机号、邮箱、部门等信息后批量导入",
    )
    add_body(
        doc,
        "管理员可在\"团队设置 > 成员管理\"中设置新成员默认角色（成员/管理员），"
        "并配置加入审批策略——可选择\"允许任何人加入\"、\"需管理员审批\"或\"仅限邀请\"。",
    )

    # 3.3 创建项目
    add_heading_styled(doc, "3.3  创建项目", level=2)
    add_body(
        doc,
        "进入\"任务管理\"模块，点击右上角\"+\"按钮选择\"新建项目\"，在弹出的创建窗口中填写项目信息：",
        indent=True,
    )
    add_bullet(doc, "项目名称：简洁明了地描述项目目标（必填）")
    add_bullet(doc, "项目描述：详细说明项目背景、目标与交付物（可选，支持富文本）")
    add_bullet(doc, "起止日期：设置项目的计划开始与结束日期，用于进度跟踪与预警")
    add_bullet(
        doc,
        "项目模板：可选空白模板、敏捷开发模板、市场营销模板或通用项目模板，模板预置了常用任务结构与字段",
    )
    add_body(
        doc,
        "项目创建后，可在项目设置中进一步配置成员权限、自定义字段、自动化规则等高级选项。"
        "项目经理可通过甘特图视图直观查看项目进度，拖拽调整任务排期。",
    )

    doc.add_page_break()

    # ========================================================================
    # 四、常见问题
    # ========================================================================
    add_heading_styled(doc, "四、常见问题 (FAQ)", level=1)
    add_body(doc, "以下汇集了用户在使用CloudWork过程中经常咨询的问题，供您参考。", indent=True)

    faqs = [
        (
            "是否支持私有化部署？",
            "支持。CloudWork为金融、政府、军工等对数据驻留要求严格的行业提供完整的私有化部署方案。"
            "私有化版本支持与客户现有AD/LDAP、SSO等系统集成，部署周期通常为5-15个工作日。"
            "如有私有化部署需求，请联系销售团队 sales@example.com 获取专属方案。",
        ),
        (
            "免费版有哪些限制？",
            "CloudWork免费版（基础版）支持最多50名成员，视频会议单场时长上限45分钟，"
            "每用户存储空间5GB。付费专业版定价为29元/人/月，包含不限时会议、100GB/人存储、"
            "高级安全功能与API接口。选择按年付费可享8折优惠，另有企业版（49元/人/月）"
            "提供专属CSM、定制品牌与SLA保障。",
        ),
        (
            "数据安全如何保障？",
            "CloudWork已通过ISO 27001认证与中国等保三级测评。技术层面，全链路采用TLS 1.3传输加密"
            "与AES-256存储加密，数据库每日自动备份并异地容灾。恢复时间目标（RTO）小于1小时，"
            "恢复点目标（RPO）小于5分钟。平台每年聘请第三方安全机构进行渗透测试，报告可向企业客户提供。",
        ),
        (
            "如何获取技术支持？",
            "标准技术支持时间为工作日9:00-18:00，可通过服务热线400-XXX-XXXX或在线客服获取帮助。"
            "紧急故障提供7×24小时值班电话（需企业客户开通紧急支持权益）。企业版客户配备专属客户成功经理"
            "（CSM），定期上门回访并提供产品培训。此外，帮助中心（help.cloudwork.example.com）"
            "提供完整的用户手册、视频教程与API文档。",
        ),
        (
            "是否支持与第三方系统集成？",
            "CloudWork提供标准RESTful API与Webhook，支持与企业现有OA、ERP、CRM等系统对接。"
            "平台已预置钉钉、企业微信、飞书等主流办公平台的单点登录（SSO）集成，"
            "同时支持SAML 2.0、OAuth 2.0、OpenID Connect等标准协议。定制化集成需求请联系解决方案架构师评估。",
        ),
    ]

    for q, a in faqs:
        add_qa(doc, q, a)

    # ========================================================================
    # FOOTER
    # ========================================================================
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(30)
    fr = footer_p.add_run("— 本文档由XX科技产品部编制，如有疑问请联系 product@example.com —")
    _set_run(fr, size=Pt(9), color=GREY)

    # ========================================================================
    # SAVE
    # ========================================================================
    out_dir = os.path.dirname(__file__)
    if not out_dir:
        out_dir = "."
    output_path = os.path.join(out_dir, "产品手册-云办公平台.docx")
    doc.save(output_path)
    size_kb = os.path.getsize(output_path) / 1024
    print(f"已生成：{output_path}")
    print(f"文件大小：{size_kb:.1f} KB")
    return output_path


if __name__ == "__main__":
    generate()
