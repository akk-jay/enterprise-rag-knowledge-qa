# -*- coding: utf-8 -*-
"""
Generate 财务报销制度.docx for XX科技有限公司 (simulated enterprise document).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

LQ = '“'  # Chinese left double quote
RQ = '”'  # Chinese right double quote
LA = '《'  # Chinese left angle bracket 《
RA = '》'  # Chinese right angle bracket 》
ID = '、'  # ideographic comma 、

def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_heading_styled(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return heading

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_bold_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table_with_data(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, '1A56DB')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'EBF0FA')
    doc.add_paragraph()
    return table

def generate():
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ======== COVER PAGE ========
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('XX科技有限公司')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('财务报销制度')
    run.font.size = Pt(36)
    run.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()
    doc.add_paragraph()

    cover_info = [
        ('版本号：', 'V4.0'),
        ('生效日期：', '2025年01月01日'),
        ('发布部门：', '财务部'),
        ('密级：', '内部公开'),
    ]
    for label, value in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run = p.add_run(value)
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ======== TOC ========
    add_heading_styled(doc, '目录', 1)
    toc_items = [
        '一、总则',
        '二、报销流程',
        '    2.1 标准报销流程',
        '    2.2 加急报销',
        '三、发票要求',
        '    3.1 发票类型',
        '    3.2 发票抬头',
        '    3.3 发票丢失处理',
        '四、费用标准',
        '    4.1 差旅费',
        '    4.2 招待费',
        '    4.3 办公用品',
        '五、审批权限',
        '六、违规处理',
        '七、附则',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.8
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ======== 一、总则 ========
    add_heading_styled(doc, '一、总则', 1)

    add_bold_body(doc, '第一条  目的')
    add_body(doc,
        '为规范XX科技有限公司（以下简称' + LQ + '公司' + RQ + '）财务报销行为，合理控制成本费用，'
        '确保资金使用合规、透明、高效，依据国家相关财税法规及公司《财务管理基本制度》，制定本制度。'
    )

    add_bold_body(doc, '第二条  适用范围')
    add_body(doc,
        '本制度适用于公司全体正式员工、试用期员工、实习生及劳务派遣人员。'
        '公司董事、监事及高级管理人员的报销行为除遵守本制度外，还需符合公司章程及相关治理文件的规定。'
    )

    add_bold_body(doc, '第三条  基本原则')
    add_body(doc,
        '公司费用报销遵循' + LQ + '事前审批、事中控制、事后报销' + RQ + '的基本原则：\n'
        '（一）事前审批：所有预计发生费用须事先通过OA系统提交申请，经有权审批人批准后方可执行。\n'
        '（二）事中控制：费用发生过程中应严格控制支出，不得超出审批金额；如需超支，须按原审批路径补办手续。\n'
        '（三）事后报销：费用实际发生后，报销人须在30个自然日内完成报销申报，逾期不予受理。'
    )

    add_bold_body(doc, '第四条  归口管理部门')
    add_body(doc,
        '财务部为本制度的归口管理部门，负责制度修订、培训宣贯、报销审核、费用分析及违规查处等工作。'
        '各部门负责人对本部门费用报销的真实性、合理性承担管理责任。'
    )

    # ======== 二、报销流程 ========
    add_heading_styled(doc, '二、报销流程', 1)

    add_heading_styled(doc, '2.1 标准报销流程', 2)
    add_body(doc,
        '公司全部费用报销均须通过OA系统' + LQ + '费用报销' + RQ + '模块在线提交，不再受理纸质报销单。'
        '标准报销流程如下：'
    )

    flow_headers = ['步骤', '环节', '责任人/角色', '时限要求', '操作说明']
    flow_rows = [
        ['①', '提交报销单', '报销人', '费用发生后30日内',
         '填写报销单并上传发票影像件（照片/扫描件），确保信息完整准确'],
        ['②', '直接上级审批', '直属管理者', '2个工作日内',
         '审核业务真实性、费用合理性；对于不合规项可退回或拒绝'],
        ['③', '部门负责人审批', '部门总监/VP', '2个工作日内',
         '单笔金额超过5000元的，还需CFO联合审批'],
        ['④', '财务审核', '财务部会计', '2个工作日内',
         '审核发票合规性、金额准确性、审批链完整性'],
        ['⑤', '出纳付款', '财务部出纳', '3个工作日内',
         '款项直接汇入报销人工资卡，原则上不支付现金'],
    ]
    add_table_with_data(doc, flow_headers, flow_rows)
    add_body(doc,
        '备注：如报销单被退回修改，修改后重新提交的，审批时限从修改提交之日起重新计算。'
    )

    add_heading_styled(doc, '2.2 加急报销', 2)
    add_body(doc,
        '因业务急需（如紧急差旅、客户接待等），报销人可在OA提交时勾选' + LQ + '加急' + RQ
        + '选项并简述加急原因。\n'
        '加急报销特殊规则：\n'
        '（一）财务部须在收到加急申请的1个工作日内完成付款。\n'
        '（二）每位员工每月加急报销不得超过2次，超出次数仍按标准流程处理。\n'
        '（三）如发现滥用加急机制（如频繁加急、虚假加急理由等），一经核实，取消该员工加急报销资格6个月，'
        '并由部门负责人对其进行严肃谈话。'
    )

    # ======== 三、发票要求 ========
    add_heading_styled(doc, '三、发票要求', 1)

    add_heading_styled(doc, '3.1 发票类型', 2)
    add_body(doc,
        '公司接受以下类型发票：\n'
        '（一）增值税专用发票（抵扣联+发票联）。\n'
        '（二）增值税普通发票（含电子普通发票）。\n'
        '（三）电子发票（PDF/OFD格式），须于国家税务总局全国增值税发票查验平台验证真伪，'
        '并将验证截图随报销附件一并上传。\n'
        '（四）定额发票（仅限单张面额不超过100元、单次报销不超过500元）。\n'
        '（五）交通票据（航空运输电子客票行程单、火车票、网约车电子发票等）。\n\n'
        '重要提示：发票开具日期须在费用实际发生之日起15个自然日内；超过15日的发票原则上不予报销，'
        '特殊情况须附书面说明并经部门负责人及财务部双签审批。'
    )

    add_heading_styled(doc, '3.2 发票抬头', 2)
    add_body(doc,
        '所有报销发票的购买方信息须严格按照以下内容开具：\n\n'
        '    单位名称：XX科技有限公司\n'
        '    纳税人识别号：91110108XXXXXXXXXX\n'
        '    地址、电话：北京市海淀区XX路XX号  010-XXXXXXXX\n'
        '    开户行及账号：招商银行北京分行XX支行  1109XXXXXXXXXXXX\n\n'
        '发票抬头信息错误（包括单位名称缺字、错字，纳税人识别号错误等）的发票不予报销，'
        '须退回开票方重新开具后重新提交。'
    )

    add_heading_styled(doc, '3.3 发票丢失处理', 2)
    add_body(doc,
        '发票丢失须在5个工作日内向财务部报告，并按以下方式处理：\n\n'
        '（一）增值税专用发票丢失：须取得开票方出具的加盖发票专用章的发票记账联复印件，'
        '同时提交书面情况说明（含丢失原因、发票信息、经办人签字）。\n\n'
        '（二）增值税普通发票/电子发票丢失：可申请开票方重新开具，或提供加盖发票专用章的发票复印件。\n\n'
        '（三）交通票据丢失（如火车票、航空行程单）：\n'
        '    1）电子行程单可登录航空公司/铁路12306官网重打或下载电子版。\n'
        '    2）无法重新获取的，报销人需如实申报，按实际发生额的80%报销'
        '（单次最高扣减不超过200元），差额部分由个人承担。\n\n'
        '（四）定额发票丢失不予补办，相关费用由个人承担。'
    )

    # ======== 四、费用标准 ========
    add_heading_styled(doc, '四、费用标准', 1)

    add_heading_styled(doc, '4.1 差旅费', 2)
    add_body(doc,
        '差旅费包括交通费、住宿费和伙食补助费三项，标准如下。员工级别以公司HR系统最新职级为准。'
    )

    add_bold_body(doc, '（1）交通工具标准')
    trans_headers = ['员工级别', '火车/高铁', '飞机', '备注']
    trans_rows = [
        ['P5及以下', '二等座', '经济舱', '动车/高铁优先，路程≤5小时原则上不乘飞机'],
        ['P6-P8', '一等座', '经济舱', '超过5小时航程可申请商务舱，需部门总监审批'],
        ['P9及以上', '商务座', '商务舱', '─'],
    ]
    add_table_with_data(doc, trans_headers, trans_rows)

    add_bold_body(doc, '（2）住宿费标准（元/晚/人）')
    hotel_headers = ['城市类别', '代表城市', 'P5及以下', 'P6-P8', 'P9及以上']
    hotel_rows = [
        ['一线城市', '北京、上海、广州、深圳', '400', '600', '1000'],
        ['其他城市', '除一线城市外的所有国内城市', '300', '450', '800'],
    ]
    add_table_with_data(doc, hotel_headers, hotel_rows)
    add_body(doc,
        '住宿费凭发票在限额内实报实销，超出限额部分由个人承担。同性两人同行须合住标准间。'
    )

    add_bold_body(doc, '（3）伙食补助费标准')
    meal_headers = ['出差类型', '补助标准', '说明']
    meal_rows = [
        ['全天出差（含过夜）', '100元/天', '凭出差审批单直接发放，无需发票'],
        ['半天出差（不含过夜）', '50元/天', '以出差目的地路程时间超过4小时认定'],
    ]
    add_table_with_data(doc, meal_headers, meal_rows)
    add_body(doc,
        '出差期间由接待单位统一安排用餐的，或已报销招待餐费的，相应餐次补助减半发放。'
    )

    add_heading_styled(doc, '4.2 招待费', 2)
    add_body(doc,
        '招待费指因业务需要招待客户、合作伙伴等外部人员所产生的餐费、礼品费、活动费等。'
    )

    add_bold_body(doc, '招待标准')
    ent_headers = ['招待对象类别', '人均限额', '审批要求', '备注']
    ent_rows = [
        ['一般客户/供应商', '≤200元/人', '部门经理审批', '同一事项一次为限'],
        ['重要客户/合作伙伴', '≤400元/人', '部门总监审批', '需提前报备客户名称及背景'],
        ['战略客户/政府接待', '无固定限额', '总经理审批', '须提交详细接待方案'],
    ]
    add_table_with_data(doc, ent_headers, ent_rows)

    add_body(doc,
        '招待费用审批补充规定：\n'
        '（一）单次招待金额超过3,000元的，须部门总监审批。\n'
        '（二）单次招待金额超过10,000元的，须总经理审批。\n'
        '（三）所有招待费须事前通过OA' + LQ + '业务招待申请' + RQ + '模块提交申请，'
        '注明招待对象、人数、事由、预计金额，审批通过后方可执行。事后报销须附招待申请单编号。'
    )

    add_heading_styled(doc, '4.3 办公用品', 2)
    add_body(doc,
        '办公用品采购实行' + LQ + '季度集中采购+零星自购' + RQ + '相结合的管理模式。'
    )

    add_bold_body(doc, '（1）集中采购')
    add_body(doc,
        '各部门于每季度末月（3月、6月、9月、12月）15日前向行政部提交下季度办公用品需求计划，'
        '由行政部统一汇总、比价、采购和发放。集中采购无需个别报销。'
    )

    add_bold_body(doc, '（2）零星自购')
    add_body(doc,
        '因工作急需且单次金额不超过500元的办公用品，员工可自行采购后凭发票报销；'
        '单次金额超过500元的，必须通过行政部统一采购，不得自行购买后报销。'
    )

    add_bold_body(doc, '（3）重要区分——非办公用品')
    add_body(doc,
        '以下物品不属于办公用品，须按固定资产或IT设备采购流程处理，不得通过办公用品科目报销：\n'
        '    · 显示器、打印机、扫描仪等电子设备\n'
        '    · 办公桌椅、文件柜等办公家具\n'
        '    · 单价超过2,000元且使用年限超过一年的物品\n'
        '上述物品须通过公司固定资产采购流程，由行政部统一采购并录入资产管理系统。'
    )

    # ======== 五、审批权限 ========
    add_heading_styled(doc, '五、审批权限', 1)

    add_body(doc,
        '各类费用报销的审批权限依据单笔金额分级设置：'
    )

    approve_headers = ['单笔金额', '审批人', '说明']
    approve_rows = [
        ['< 2,000元', '直接上级', '差旅、招待费至少2级审批'],
        ['2,000元 - 5,000元', '直接上级 → 部门总监', '─'],
        ['5,000元 - 20,000元', '直接上级 → 部门总监 → CFO', '─'],
        ['> 20,000元', '直接上级 → 部门总监 → CFO → 总经理', '─'],
    ]
    add_table_with_data(doc, approve_headers, approve_rows)

    add_bold_body(doc, '特别规定：')
    add_body(doc,
        '（一）差旅费和招待费无论金额大小，须至少经过两级审批（直接上级+部门总监），不得仅由一级审批。\n'
        '（二）审批人不得审批自己的费用报销，须自动提交至上一级审批人。\n'
        '（三）如审批人外出超过3个工作日，应事先在OA系统中设置代理审批人，避免审批积压。\n'
        '（四）涉及多个部门共同发生的费用，由主要承担部门提交报销，其他部门会签确认。'
    )

    # ======== 六、违规处理 ========
    add_heading_styled(doc, '六、违规处理', 1)

    add_bold_body(doc, '第一条  虚假报销认定')
    add_body(doc,
        '以下行为属于虚假报销：\n'
        '（一）伪造、变造、涂改发票或费用凭证。\n'
        '（二）同一笔费用重复提交报销。\n'
        '（三）个人消费冒充公务消费报销。\n'
        '（四）虚增报销金额（如虚报人数、天数、单价等）。\n'
        '（五）冒用他人名义报销。'
    )

    add_bold_body(doc, '第二条  处罚措施')
    add_body(doc,
        '一经查实存在虚假报销行为的：\n'
        '（一）本次报销申请立即作废，已付款项须全额退还公司。\n'
        '（二）处以虚报金额2倍的罚款（最低罚款金额500元）。\n'
        '（三）在公司内部通报批评，记入个人诚信档案。\n'
        '（四）情节严重或累计两次及以上者，公司有权依据《劳动合同法》第三十九条解除劳动合同，'
        '且不支付经济补偿金。涉嫌违法的，移送司法机关处理。'
    )

    add_bold_body(doc, '第三条  审批人连带责任')
    add_body(doc,
        '审批人对其审批通过的报销事项承担连带审核责任。若因审批人未尽合理审核义务导致虚假报销通过，'
        '将对审批人进行绩效扣分（每次扣5分）并视情节追究管理责任。'
    )

    add_bold_body(doc, '第四条  财务抽查制度')
    add_body(doc,
        '财务部每季度按不少于当季报销申请总笔数5%的比例进行随机抽查。抽查内容包括：\n'
        '（一）发票真伪验证；（二）费用实际发生情况核查；（三）审批流程规范性检查。\n'
        '抽查结果形成书面报告，抄送公司管理层。'
    )

    add_bold_body(doc, '第五条  举报与奖励')
    add_body(doc,
        '公司鼓励员工对违规报销行为进行监督举报。举报渠道：检举邮箱 jiancha@公司域名.com'
        '（财务部负责人和内审负责人双人接收）。经核实属实的，给予举报人500元至5,000元的奖励。'
        '公司严格保护举报人信息，禁止任何形式的打击报复。'
    )

    # ======== 七、附则 ========
    add_heading_styled(doc, '七、附则', 1)

    add_body(doc,
        '（一）本制度由财务部负责解释和修订。\n\n'
        '（二）本制度自2025年01月01日起施行，原《财务报销管理暂行办法》（V3.0）同时废止。\n\n'
        '（三）本制度未尽事宜，按照国家有关法律法规及公司其他相关制度执行。\n\n'
        '（四）各部门可依据本制度制定部门级实施细则，但不得与本制度相抵触，制定前须报财务部备案。'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('XX科技有限公司  财务部')
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('2025年01月01日')
    run.font.size = Pt(10.5)

    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, '财务报销制度.docx')
    doc.save(output_path)
    print('已生成：' + output_path)
    return output_path

if __name__ == '__main__':
    generate()
