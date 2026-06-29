"""
parser.py — 文档解析器

支持三种格式：
  - PDF  → 用 PyMuPDF (fitz) 提取文本
  - DOCX → 用 python-docx 提取段落和表格
  - TXT  → 直接读取（UTF-8 / GBK）

核心函数: parse_document(file_path) -> dict
"""
import re
from pathlib import Path


def parse_document(file_path):
    """
    解析文档，返回纯文本和元信息。

    参数:
        file_path: 文档文件的路径（字符串或 Path 对象）

    返回:
        dict {
            "text": "文档的完整纯文本内容",
            "filename": "文档文件名",
            "format": "pdf" | "docx" | "txt"
        }

    抛出:
        ValueError: 不支持的文件格式
        FileNotFoundError: 文件不存在
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    # 根据扩展名判断格式（不区分大小写）
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        text = _parse_pdf(file_path)
        fmt = "pdf"
    elif ext == ".docx":
        text = _parse_docx(file_path)
        fmt = "docx"
    elif ext == ".txt":
        text = _parse_txt(file_path)
        fmt = "txt"
    else:
        raise ValueError(
            f"不支持的文档格式: {ext}。"
            f"当前支持: PDF (.pdf), Word (.docx), 纯文本 (.txt)"
        )

    # 清理多余空白
    text = _clean_text(text)

    return {
        "text": text,
        "filename": file_path.name,
        "format": fmt,
    }


def _parse_pdf(file_path):
    """用 PyMuPDF 逐页提取 PDF 文本"""
    import fitz  # PyMuPDF 的导入名

    doc = fitz.open(str(file_path))
    pages = []

    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text("text")
            if page_text.strip():
                pages.append(page_text)
    finally:
        doc.close()

    return "\n".join(pages)


def _parse_docx(file_path):
    """用 python-docx 提取段落和表格文本"""
    from docx import Document

    doc = Document(str(file_path))
    parts = []

    # 1. 提取段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 2. 提取表格（财务报销制度等文档的重要内容在表格里）
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _parse_txt(file_path):
    """读取纯文本文件，UTF-8 失败则尝试 GBK"""
    # 先尝试 UTF-8（最常用的中文编码）
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        pass

    # 再尝试 GBK（Windows 中文默认编码）
    try:
        with open(file_path, "r", encoding="gbk") as f:
            return f.read()
    except UnicodeDecodeError:
        pass

    # 最后尝试系统默认编码
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _clean_text(text):
    """清理文本中的多余空白"""
    # 3 个以上连续换行 → 2 个换行
    text = re.sub(r"\n{3,}", "\n\n", text)
    # 去掉行首行尾多余空格（保留有意义的内容）
    text = re.sub(r"[ \t]+", " ", text)
    # 去掉空行前后的空白
    text = text.strip()
    return text


# ── 快速测试 ──
if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        result = parse_document(sys.argv[1])
        print(f"文件名: {result['filename']}")
        print(f"格式: {result['format']}")
        print(f"字符数: {len(result['text'])}")
        print(f"\n--- 前 500 字符预览 ---")
        print(result["text"][:500])
    else:
        print("用法: python parser.py <文件路径>")
